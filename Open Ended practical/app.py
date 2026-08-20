"""
Document Vision Application Dashboard (FastAPI Backend + Modern Glassmorphism Web UI).

Provides an end-to-end web portal for government offices to convert scanned
documents into editable text with live OpenCV preprocessing controls, OCR extraction,
entity validation, and export tools.
Made by Yash Kapse.
"""

import os
import sys
import io
import base64
import json
import cv2
import numpy as np
import traceback
from typing import Optional
from pydantic import BaseModel
from fastapi import FastAPI, Request
from fastapi.responses import HTMLResponse, JSONResponse, StreamingResponse
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
import docx
from docx import Document

# Ensure core modules are importable
sys.path.append(os.path.abspath(os.path.dirname(__file__)))
from core.preprocessor import DocumentPreprocessor
from core.ocr_engine import OCREngine
from core.validator import DocumentValidator
from samples.synthetic_docs import create_synthetic_government_doc

app = FastAPI(
    title="Government Document Vision & OCR Application — Made by Yash Kapse",
    description="Convert scanned government documents to editable text with OpenCV & OCR validation. Made by Yash Kapse.",
    version="1.0.0"
)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

preprocessor = DocumentPreprocessor()
ocr_engine = OCREngine()
validator = DocumentValidator()


class ProcessRequest(BaseModel):
    image_base64: Optional[str] = None
    ground_truth: Optional[str] = None
    auto_deskew: bool = True
    denoise_method: str = "gaussian"
    denoise_kernel: int = 5
    threshold_method: str = "otsu"
    adaptive_block_size: int = 15
    adaptive_c: int = 4
    apply_morphology: bool = False
    unwarp: bool = False
    engine: str = "auto"


def mat_to_base64(img: np.ndarray, format: str = ".png") -> str:
    """Helper to convert OpenCV image matrix to base64 string."""
    if img is None:
        return ""
    success, buffer = cv2.imencode(format, img)
    if not success:
        return ""
    encoded = base64.b64encode(buffer).decode("utf-8")
    return f"data:image/png;base64,{encoded}"


@app.get("/api/generate_sample")
async def generate_sample_document(skew: float = 4.5, noise: float = 10.0):
    """Generate sample synthetic government document for testing."""
    try:
        img, gt_text, fields = create_synthetic_government_doc(skew_angle=skew, noise_level=noise)
        img_b64 = mat_to_base64(img)
        return {
            "success": True,
            "image_base64": img_b64,
            "ground_truth": gt_text,
            "fields": fields
        }
    except Exception as e:
        return JSONResponse(status_code=500, content={"success": False, "error": str(e)})


@app.post("/api/process")
async def process_document(req: ProcessRequest):
    """
    Process document image through preprocessing, OCR extraction, and validation.
    Receives JSON payload to bypass Starlette multipart file size limits.
    """
    try:
        if not req.image_base64:
            return JSONResponse(
                status_code=400,
                content={"success": False, "error": "No image data provided."}
            )

        clean_b64 = req.image_base64.strip()
        if "," in clean_b64:
            clean_b64 = clean_b64.split(",")[1]

        img_bytes = base64.b64decode(clean_b64)
        nparr = np.frombuffer(img_bytes, np.uint8)
        raw_img = cv2.imdecode(nparr, cv2.IMREAD_COLOR)

        if raw_img is None:
            return JSONResponse(
                status_code=400,
                content={"success": False, "error": "Invalid image format. Could not decode image."}
            )

        # Downscale ultra-high-resolution images (>2500px) slightly for speed & memory
        h, w = raw_img.shape[:2]
        if max(h, w) > 2500:
            scale = 2500.0 / max(h, w)
            raw_img = cv2.resize(raw_img, (int(w * scale), int(h * scale)), interpolation=cv2.INTER_AREA)

        # 2. Run Preprocessing Pipeline
        config = {
            "auto_deskew": req.auto_deskew,
            "denoise_method": req.denoise_method,
            "denoise_kernel": req.denoise_kernel,
            "threshold_method": req.threshold_method,
            "adaptive_block_size": req.adaptive_block_size,
            "adaptive_c": req.adaptive_c,
            "apply_morphology": req.apply_morphology,
            "unwarp": req.unwarp
        }

        prep_res = preprocessor.process(raw_img, config=config)

        # 3. Run OCR Text Extraction
        ocr_res = ocr_engine.extract_text(prep_res["final"], engine=req.engine)

        # 4. Generate Bounding Box Overlay Image
        bbox_vis = ocr_engine.visualize_boxes(prep_res["gray"], ocr_res["words"])

        # 5. Entity Extraction & Validation
        val_res = validator.validate_document(ocr_res["full_text"], ground_truth=req.ground_truth)

        return {
            "success": True,
            "skew_angle": round(prep_res["skew_angle"], 2),
            "ocr_engine": ocr_res["engine"],
            "total_words": ocr_res["total_words"],
            "avg_confidence": round(ocr_res["avg_confidence"] * 100, 1),
            "extracted_text": ocr_res["full_text"],
            "validation": val_res,
            "images": {
                "original": mat_to_base64(prep_res["original"]),
                "gray": mat_to_base64(prep_res["gray"]),
                "denoised": mat_to_base64(prep_res["denoised"]),
                "binary": mat_to_base64(prep_res["binary"]),
                "bbox_overlay": mat_to_base64(bbox_vis)
            }
        }
    except Exception as e:
        traceback.print_exc()
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": f"Processing Error: {str(e)}"}
        )


@app.post("/api/export")
async def export_document(request: Request):
    """Export extracted editable text to TXT, DOCX, or JSON format."""
    try:
        data = await request.json()
        text = str(data.get("text", ""))
        fmt = str(data.get("format", "txt")).lower()

        if fmt == "txt":
            stream = io.BytesIO(text.encode("utf-8"))
            return StreamingResponse(
                stream,
                media_type="text/plain",
                headers={"Content-Disposition": "attachment; filename=extracted_document.txt"}
            )
        elif fmt == "json":
            json_data = {"extracted_text": text, "author": "Made by Yash Kapse"}
            stream = io.BytesIO(json.dumps(json_data, indent=2).encode("utf-8"))
            return StreamingResponse(
                stream,
                media_type="application/json",
                headers={"Content-Disposition": "attachment; filename=extracted_document.json"}
            )
        elif fmt == "docx":
            doc = Document()
            doc.add_heading('Converted Government Document Text', 0)
            doc.add_paragraph('Made by Yash Kapse — Document Vision Application')
            doc.add_paragraph('--------------------------------------------------')
            for line in text.splitlines():
                doc.add_paragraph(line)
            stream = io.BytesIO()
            doc.save(stream)
            stream.seek(0)
            return StreamingResponse(
                stream,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": "attachment; filename=extracted_document.docx"}
            )
        return JSONResponse(status_code=400, content={"error": "Unsupported export format."})
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})


@app.get("/", response_class=HTMLResponse)
async def serve_ui():
    """Render single-page Document Vision application frontend (White & Olive Theme)."""
    return """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>GovVision OCR — Made by Yash Kapse</title>
    <link href="https://fonts.googleapis.com/css2?family=Outfit:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap" rel="stylesheet">
    <style>
        :root {
            --bg-gradient: linear-gradient(135deg, #f8fafc 0%, #f1f5f9 50%, #e2e8f0 100%);
            --panel-bg: #ffffff;
            --panel-border: rgba(85, 107, 47, 0.2);
            --panel-shadow: 0 4px 20px rgba(85, 107, 47, 0.08);
            
            --olive-primary: #556b2f;
            --olive-dark: #3e4e22;
            --olive-light: #6b8e23;
            --olive-soft: rgba(85, 107, 47, 0.08);
            --olive-hover: #485c27;
            
            --text-main: #1e293b;
            --text-muted: #64748b;
            --card-border: #e2e8f0;
        }

        * { box-sizing: border-box; margin: 0; padding: 0; }

        body {
            font-family: 'Outfit', sans-serif;
            background: var(--bg-gradient);
            color: var(--text-main);
            min-height: 100vh;
            display: flex;
            flex-direction: column;
        }

        header {
            background: #ffffff;
            border-bottom: 2px solid var(--olive-primary);
            padding: 1rem 2.5rem;
            display: flex;
            justify-content: space-between;
            align-items: center;
            box-shadow: 0 2px 10px rgba(85, 107, 47, 0.1);
        }

        .logo-container {
            display: flex;
            align-items: center;
            gap: 14px;
        }

        .logo-icon {
            width: 44px;
            height: 44px;
            background: linear-gradient(135deg, #556b2f, #6b8e23);
            border-radius: 12px;
            display: flex;
            align-items: center;
            justify-content: center;
            font-weight: 800;
            font-size: 20px;
            color: #ffffff;
            box-shadow: 0 4px 12px rgba(85, 107, 47, 0.3);
        }

        .app-title { font-size: 1.35rem; font-weight: 700; color: #1e293b; letter-spacing: -0.5px; }
        .app-subtitle { font-size: 0.82rem; color: var(--text-muted); }

        .author-badge {
            background: linear-gradient(135deg, #556b2f, #6b8e23);
            color: #ffffff;
            padding: 6px 16px;
            border-radius: 20px;
            font-size: 0.85rem;
            font-weight: 600;
            letter-spacing: 0.3px;
            box-shadow: 0 2px 8px rgba(85, 107, 47, 0.25);
            display: inline-flex;
            align-items: center;
            gap: 6px;
        }

        .btn {
            background: var(--olive-primary);
            color: white;
            border: none;
            padding: 10px 22px;
            border-radius: 8px;
            font-family: inherit;
            font-weight: 600;
            font-size: 0.9rem;
            cursor: pointer;
            transition: all 0.2s ease;
            display: inline-flex;
            align-items: center;
            gap: 8px;
            box-shadow: 0 3px 10px rgba(85, 107, 47, 0.25);
        }
        .btn:hover { background: var(--olive-hover); transform: translateY(-1px); box-shadow: 0 5px 15px rgba(85, 107, 47, 0.35); }
        .btn-secondary { background: #f1f5f9; color: var(--olive-dark); border: 1px solid var(--panel-border); box-shadow: none; }
        .btn-secondary:hover { background: #e2e8f0; color: var(--olive-primary); }
        .btn:disabled { opacity: 0.6; cursor: not-allowed; }

        main {
            flex: 1;
            padding: 1.5rem 2.5rem;
            display: grid;
            grid-template-columns: 340px 1fr;
            gap: 1.5rem;
            max-width: 1650px;
            margin: 0 auto;
            width: 100%;
        }

        .sidebar {
            background: var(--panel-bg);
            border: 1px solid var(--panel-border);
            border-radius: 16px;
            padding: 1.4rem;
            display: flex;
            flex-direction: column;
            gap: 1.4rem;
            box-shadow: var(--panel-shadow);
        }

        .section-title {
            font-size: 0.9rem;
            font-weight: 700;
            text-transform: uppercase;
            letter-spacing: 1px;
            color: var(--olive-primary);
            margin-bottom: 0.75rem;
            display: flex;
            align-items: center;
            justify-content: space-between;
        }

        .drop-zone {
            border: 2px dashed var(--olive-primary);
            border-radius: 12px;
            padding: 1.5rem;
            text-align: center;
            background: var(--olive-soft);
            cursor: pointer;
            transition: all 0.2s ease;
        }
        .drop-zone:hover { border-color: var(--olive-dark); background: rgba(85, 107, 47, 0.14); }

        .control-group {
            display: flex;
            flex-direction: column;
            gap: 8px;
            margin-bottom: 12px;
        }
        .control-label { font-size: 0.85rem; font-weight: 600; color: var(--text-main); display: flex; justify-content: space-between; }
        select {
            width: 100%;
            background: #f8fafc;
            border: 1px solid var(--panel-border);
            color: var(--text-main);
            padding: 9px 12px;
            border-radius: 8px;
            font-family: inherit;
            font-size: 0.88rem;
            outline: none;
            transition: border-color 0.2s;
        }
        select:focus { border-color: var(--olive-primary); }

        .toggle-switch {
            display: flex;
            align-items: center;
            justify-content: space-between;
            font-size: 0.85rem;
            font-weight: 500;
            padding: 6px 0;
            color: var(--text-main);
        }

        .content-area {
            display: flex;
            flex-direction: column;
            gap: 1.5rem;
        }

        .metrics-row {
            display: grid;
            grid-template-columns: repeat(4, 1fr);
            gap: 1rem;
        }

        .metric-card {
            background: var(--panel-bg);
            border: 1px solid var(--panel-border);
            border-radius: 14px;
            padding: 1.1rem 1.25rem;
            display: flex;
            flex-direction: column;
            gap: 4px;
            box-shadow: var(--panel-shadow);
        }
        .metric-val { font-size: 1.6rem; font-weight: 800; color: var(--olive-primary); }
        .metric-lbl { font-size: 0.75rem; color: var(--text-muted); font-weight: 600; text-transform: uppercase; letter-spacing: 0.5px; }

        .workspace-grid {
            display: grid;
            grid-template-columns: 1fr 1fr;
            gap: 1.5rem;
            flex: 1;
        }

        .panel {
            background: var(--panel-bg);
            border: 1px solid var(--panel-border);
            border-radius: 16px;
            padding: 1.4rem;
            display: flex;
            flex-direction: column;
            gap: 1rem;
            box-shadow: var(--panel-shadow);
        }

        .tabs {
            display: flex;
            gap: 8px;
            border-bottom: 2px solid #f1f5f9;
            padding-bottom: 8px;
        }
        .tab {
            padding: 7px 16px;
            border-radius: 8px;
            font-size: 0.85rem;
            font-weight: 600;
            cursor: pointer;
            color: var(--text-muted);
            transition: all 0.2s;
        }
        .tab.active { background: var(--olive-primary); color: white; }
        .tab:hover:not(.active) { background: #f1f5f9; color: var(--olive-primary); }

        .img-preview-container {
            flex: 1;
            min-height: 400px;
            background: #f8fafc;
            border-radius: 12px;
            display: flex;
            align-items: center;
            justify-content: center;
            overflow: hidden;
            border: 1px solid #e2e8f0;
            position: relative;
        }
        .img-preview-container img { max-width: 100%; max-height: 540px; object-fit: contain; }

        textarea {
            width: 100%;
            height: 100%;
            min-height: 400px;
            background: #fdfdfd;
            border: 1px solid #cbd5e1;
            color: #0f172a;
            padding: 1.1rem;
            border-radius: 12px;
            font-family: 'JetBrains Mono', monospace;
            font-size: 0.9rem;
            line-height: 1.6;
            resize: vertical;
            outline: none;
        }
        textarea:focus { border-color: var(--olive-primary); box-shadow: 0 0 0 3px rgba(85, 107, 47, 0.15); }

        .entity-chip {
            background: var(--olive-soft);
            border: 1px solid rgba(85, 107, 47, 0.3);
            color: var(--olive-dark);
            padding: 7px 14px;
            border-radius: 8px;
            font-size: 0.82rem;
            font-weight: 500;
            display: inline-flex;
            gap: 6px;
            align-items: center;
            margin: 4px;
        }

        footer {
            background: #ffffff;
            border-top: 1px solid var(--panel-border);
            padding: 0.8rem 2.5rem;
            text-align: center;
            font-size: 0.85rem;
            color: var(--text-muted);
            font-weight: 500;
        }
        footer span { color: var(--olive-primary); font-weight: 700; }
    </style>
</head>
<body>
    <header>
        <div class="logo-container">
            <div class="logo-icon">GV</div>
            <div>
                <div class="app-title">GovVision OCR Portal</div>
                <div class="app-subtitle">Document Preprocessing & Automated OCR Vision Application</div>
            </div>
        </div>
        <div style="display: flex; align-items: center; gap: 16px;">
            <div class="author-badge">👤 Made by Yash Kapse</div>
            <button class="btn btn-secondary" onclick="generateSampleDoc()">✨ Generate Sample Doc</button>
            <button class="btn" id="btnProcess" onclick="triggerProcess()">🚀 Run OCR Processing</button>
        </div>
    </header>

    <main>
        <div class="sidebar">
            <div>
                <div class="section-title">📄 Document Input</div>
                <div class="drop-zone" id="dropZone" onclick="document.getElementById('fileInput').click()">
                    <span style="font-size: 32px;">📤</span>
                    <div style="margin-top: 8px; font-weight: 600; font-size: 0.92rem;" id="fileNameDisplay">Click or Drag Document</div>
                    <div style="font-size: 0.78rem; color: var(--text-muted); margin-top: 4px;">PNG, JPG, TIFF (Scanned Forms)</div>
                    <input type="file" id="fileInput" accept="image/*" style="display: none;" onchange="handleFileSelect(event)">
                </div>
            </div>

            <div>
                <div class="section-title">⚙️ Preprocessing Controls</div>

                <div class="toggle-switch">
                    <span>Auto Deskewing</span>
                    <input type="checkbox" id="autoDeskew" checked>
                </div>

                <div class="control-group">
                    <div class="control-label">Denoise Method</div>
                    <select id="denoiseMethod">
                        <option value="gaussian">Gaussian Filter</option>
                        <option value="median">Median Filter</option>
                        <option value="bilateral">Bilateral Filter</option>
                    </select>
                </div>

                <div class="control-group">
                    <div class="control-label">Threshold Binarization</div>
                    <select id="thresholdMethod">
                        <option value="otsu">Otsu Thresholding</option>
                        <option value="adaptive_gaussian">Adaptive Gaussian</option>
                        <option value="adaptive_mean">Adaptive Mean</option>
                    </select>
                </div>

                <div class="toggle-switch">
                    <span>Morphological Closing</span>
                    <input type="checkbox" id="applyMorphology">
                </div>
                
                <div class="toggle-switch">
                    <span>Unwarp Perspective</span>
                    <input type="checkbox" id="unwarp">
                </div>
            </div>

            <div>
                <div class="section-title">🔍 OCR Engine</div>
                <select id="ocrEngine">
                    <option value="auto">Auto Select (EasyOCR / Tesseract)</option>
                    <option value="easyocr">EasyOCR Deep Learning</option>
                    <option value="tesseract">Tesseract OCR</option>
                </select>
            </div>
        </div>

        <div class="content-area">
            <div class="metrics-row">
                <div class="metric-card">
                    <div class="metric-lbl">Detected Skew</div>
                    <div class="metric-val" id="metricSkew">0.0°</div>
                </div>
                <div class="metric-card">
                    <div class="metric-lbl">Total Words</div>
                    <div class="metric-val" id="metricWords">0</div>
                </div>
                <div class="metric-card">
                    <div class="metric-lbl">OCR Confidence</div>
                    <div class="metric-val" id="metricConf">0%</div>
                </div>
                <div class="metric-card">
                    <div class="metric-lbl">CER / WER</div>
                    <div class="metric-val" id="metricError">--</div>
                </div>
            </div>

            <div class="workspace-grid">
                <!-- Left Panel: Image Visualizer -->
                <div class="panel">
                    <div class="tabs">
                        <div class="tab active" onclick="switchImageTab('bbox_overlay', this)">OCR Bounding Boxes</div>
                        <div class="tab" onclick="switchImageTab('original', this)">Original</div>
                        <div class="tab" onclick="switchImageTab('binary', this)">Binarized</div>
                        <div class="tab" onclick="switchImageTab('denoised', this)">Denoised</div>
                    </div>
                    <div class="img-preview-container">
                        <div id="imagePlaceholder" style="color: var(--text-muted); font-size: 0.9rem;">No image loaded</div>
                        <img id="mainImagePreview" style="display: none;" src="" alt="Document Preview">
                    </div>
                </div>

                <!-- Right Panel: Text Editor & Extracted Entities -->
                <div class="panel">
                    <div style="display: flex; justify-content: space-between; align-items: center;">
                        <div class="tabs" style="border: none; padding: 0;">
                            <div class="tab active" onclick="switchRightTab('editor', this)">Editable Text</div>
                            <div class="tab" onclick="switchRightTab('entities', this)">Extracted Entities</div>
                        </div>
                        <div style="display: flex; gap: 8px;">
                            <button class="btn btn-secondary" style="padding: 5px 12px; font-size: 0.82rem;" onclick="exportText('txt')">TXT</button>
                            <button class="btn btn-secondary" style="padding: 5px 12px; font-size: 0.82rem;" onclick="exportText('docx')">DOCX</button>
                            <button class="btn btn-secondary" style="padding: 5px 12px; font-size: 0.82rem;" onclick="exportText('json')">JSON</button>
                        </div>
                    </div>

                    <div id="editorTab" style="flex: 1; display: flex; flex-direction: column;">
                        <textarea id="extractedTextArea" placeholder="Extracted document text will appear here. You can edit, format, or copy the content directly..."></textarea>
                    </div>

                    <div id="entitiesTab" style="display: none; flex: 1; flex-direction: column; gap: 12px;">
                        <div style="font-size: 0.85rem; font-weight: 700; color: var(--olive-primary);">GOVERNMENT RECOGNIZED ENTITIES</div>
                        <div id="entityChipsContainer" style="display: flex; flex-wrap: wrap;">
                            <span style="font-size: 0.85rem; color: var(--text-muted);">No entities extracted yet.</span>
                        </div>
                        
                        <div style="font-size: 0.85rem; font-weight: 700; color: var(--olive-primary); margin-top: 10px;">STRUCTURED KEY-VALUE PAIRS</div>
                        <div id="keyValueContainer" style="background: #f8fafc; padding: 14px; border-radius: 10px; border: 1px solid var(--panel-border); font-family: 'JetBrains Mono', monospace; font-size: 0.85rem; min-height: 150px;">
                        </div>
                    </div>
                </div>
            </div>
        </div>
    </main>

    <footer>
        Practical 9 — Document Vision Application | <span>Made by Yash Kapse</span>
    </footer>

    <script>
        let currentProcessData = null;
        let selectedImageBase64 = null;
        let groundTruthText = null;

        function handleFileSelect(evt) {
            const file = evt.target.files[0];
            if (!file) return;
            selectedImageBase64 = null;
            groundTruthText = null;

            document.getElementById('fileNameDisplay').innerText = file.name;

            const reader = new FileReader();
            reader.onload = function(e) {
                selectedImageBase64 = e.target.result;
                document.getElementById('mainImagePreview').src = selectedImageBase64;
                document.getElementById('mainImagePreview').style.display = 'block';
                document.getElementById('imagePlaceholder').style.display = 'none';
                triggerProcess();
            };
            reader.readAsDataURL(file);
        }

        async function generateSampleDoc() {
            try {
                document.getElementById('fileNameDisplay').innerText = 'Sample Government Document';
                const res = await fetch('/api/generate_sample?skew=4.5&noise=10.0');
                const data = await res.json();

                if (!data.success) {
                    alert("Failed to generate sample: " + data.error);
                    return;
                }

                selectedImageBase64 = data.image_base64;
                groundTruthText = data.ground_truth;
                document.getElementById('mainImagePreview').src = selectedImageBase64;
                document.getElementById('mainImagePreview').style.display = 'block';
                document.getElementById('imagePlaceholder').style.display = 'none';
                triggerProcess();
            } catch (err) {
                alert("Failed to generate sample document: " + err);
            }
        }

        async function triggerProcess() {
            if (!selectedImageBase64) {
                alert("Please upload or generate a scanned document first.");
                return;
            }

            const btn = document.getElementById('btnProcess');
            btn.disabled = true;
            btn.innerText = '⏳ Processing OCR...';

            const payload = {
                image_base64: selectedImageBase64,
                ground_truth: groundTruthText || null,
                auto_deskew: document.getElementById('autoDeskew').checked,
                denoise_method: document.getElementById('denoiseMethod').value,
                threshold_method: document.getElementById('thresholdMethod').value,
                apply_morphology: document.getElementById('applyMorphology').checked,
                unwarp: document.getElementById('unwarp').checked,
                engine: document.getElementById('ocrEngine').value
            };

            try {
                const res = await fetch('/api/process', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify(payload)
                });
                const data = await res.json();

                if (!res.ok || !data.success) {
                    alert("OCR Error: " + (data.error || "Processing failed"));
                    return;
                }

                currentProcessData = data;
                updateUI(data);
            } catch (err) {
                alert("OCR Processing Connection Error: " + err);
            } finally {
                btn.disabled = false;
                btn.innerText = '🚀 Run OCR Processing';
            }
        }

        function updateUI(data) {
            document.getElementById('metricSkew').innerText = data.skew_angle + '°';
            document.getElementById('metricWords').innerText = data.total_words;
            document.getElementById('metricConf').innerText = data.avg_confidence + '%';
            
            if (data.validation && data.validation.cer !== null) {
                document.getElementById('metricError').innerText = 'CER ' + (data.validation.cer * 100).toFixed(1) + '%';
            } else {
                document.getElementById('metricError').innerText = '--';
            }

            if (data.images && data.images.bbox_overlay) {
                document.getElementById('mainImagePreview').src = data.images.bbox_overlay;
            }
            document.getElementById('extractedTextArea').value = data.extracted_text || "";

            // Render Entity Chips
            const chipContainer = document.getElementById('entityChipsContainer');
            chipContainer.innerHTML = '';
            const entities = data.validation ? data.validation.entities : {};
            let foundCount = 0;

            for (const [key, matches] of Object.entries(entities)) {
                matches.forEach(m => {
                    foundCount++;
                    const chip = document.createElement('div');
                    chip.className = 'entity-chip';
                    chip.innerHTML = `<strong>${key.toUpperCase()}:</strong> ${m}`;
                    chipContainer.appendChild(chip);
                });
            }
            if (foundCount === 0) {
                chipContainer.innerHTML = '<span style="font-size: 0.85rem; color: var(--text-muted);">No standard patterns matched.</span>';
            }

            // Render Key-Value pairs
            const kvContainer = document.getElementById('keyValueContainer');
            const kvs = data.validation ? data.validation.key_values : {};
            if (Object.keys(kvs).length > 0) {
                let html = '<table style="width: 100%; border-collapse: collapse;">';
                for (const [k, v] of Object.entries(kvs)) {
                    html += `<tr><td style="color: var(--olive-primary); font-weight:600; padding: 6px 10px; width: 40%; border-bottom: 1px solid #e2e8f0;">${k}</td><td style="padding: 6px 10px; color: #1e293b; border-bottom: 1px solid #e2e8f0;">${v}</td></tr>`;
                }
                html += '</table>';
                kvContainer.innerHTML = html;
            } else {
                kvContainer.innerHTML = '<span style="color: var(--text-muted);">No key-value pairs detected.</span>';
            }
        }

        function switchImageTab(imageKey, el) {
            document.querySelectorAll('.panel:first-child .tab').forEach(t => t.classList.remove('active'));
            el.classList.add('active');
            if (currentProcessData && currentProcessData.images && currentProcessData.images[imageKey]) {
                document.getElementById('mainImagePreview').src = currentProcessData.images[imageKey];
            }
        }

        function switchRightTab(tabName, el) {
            document.querySelectorAll('.panel:last-child .tab').forEach(t => t.classList.remove('active'));
            el.classList.add('active');
            if (tabName === 'editor') {
                document.getElementById('editorTab').style.display = 'flex';
                document.getElementById('entitiesTab').style.display = 'none';
            } else {
                document.getElementById('editorTab').style.display = 'none';
                document.getElementById('entitiesTab').style.display = 'flex';
            }
        }

        async function exportText(format) {
            const text = document.getElementById('extractedTextArea').value;
            if (!text) {
                alert("No text to export.");
                return;
            }

            try {
                const res = await fetch('/api/export', {
                    method: 'POST',
                    headers: { 'Content-Type': 'application/json' },
                    body: JSON.stringify({ text: text, format: format })
                });
                if (!res.ok) {
                    alert("Export failed.");
                    return;
                }
                const blob = await res.blob();
                const url = window.URL.createObjectURL(blob);
                const a = document.createElement('a');
                a.href = url;
                a.download = `extracted_document.${format}`;
                document.body.appendChild(a);
                a.click();
                a.remove();
            } catch (err) {
                alert("Export error: " + err);
            }
        }

        // Auto load sample on start
        window.addEventListener('load', () => {
            generateSampleDoc();
        });
    </script>
</body>
</html>
"""


if __name__ == "__main__":
    uvicorn.run(app, host="127.0.0.1", port=8000)
